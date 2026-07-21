from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"D:\coding-document\shokz\imu-earphone")
SOURCE = ROOT / "采集计划.md"
OUTPUT = ROOT / "IMU数据采集计划_修订标红版.docx"

BLACK = RGBColor(0, 0, 0)
RED = RGBColor(192, 0, 0)
BLUE = RGBColor(46, 116, 181)
LIGHT_BLUE = "E8EEF5"


def set_run_font(run, name="Microsoft YaHei", size=10.5, color=BLACK, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=110, bottom=70, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_text(p, text, color, bold=False, size=10.5):
    # Preserve inline markdown bold as Word bold.
    tokens = re.split(r"(\*\*.*?\*\*)", text)
    for token in tokens:
        if not token:
            continue
        is_bold = token.startswith("**") and token.endswith("**")
        value = token[2:-2] if is_bold else token
        value = re.sub(r"`([^`]*)`", r"\1", value)
        run = p.add_run(value)
        set_run_font(run, color=color, bold=(bold or is_bold), size=size)


def add_heading(doc, level, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 9}.get(level, 8))
    p.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}.get(level, 4))
    p.paragraph_format.keep_with_next = True
    add_text(p, text, color, bold=True, size={1: 16, 2: 13, 3: 11.5}.get(level, 11))
    return p


def pipe_cells(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def table_widths(cols):
    if cols == 2:
        return [2700, 6660]
    if cols == 3:
        return [2300, 4200, 2860]
    if cols == 4:
        return [1800, 2700, 3000, 1860]
    base = 9360 // cols
    return [base] * (cols - 1) + [9360 - base * (cols - 1)]


def add_table(doc, lines, color):
    rows = [pipe_cells(line) for line in lines if not re.match(r"^\|?\s*:?-{3,}", line.strip())]
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_geometry(table, table_widths(cols))
    for i, values in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            if i == 0:
                shade(cell, LIGHT_BLUE)
            add_text(p, values[j] if j < len(values) else "", color, bold=(i == 0), size=9.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(header, "IMU 数据采集计划 | 修订标红版", RGBColor(100, 100, 100), size=8.5)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "红色文字为本次新增或实质调整内容", RED, size=8.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    add_text(title, "IMU 数据采集计划", BLACK, bold=True, size=22)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(16)
    add_text(sub, "耳机 IMU 姿势提醒与基础动作识别 | 修订标红版", RGBColor(90, 90, 90), size=10.5)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    add_text(note, "说明：红色文字为本次新增或实质调整内容；黑色文字为原计划保留内容。", RED, bold=True, size=10)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    active_levels = []
    active_red = False
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("#"):
            m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if m:
                level = len(m.group(1))
                text = m.group(2)
                is_marked = "【新增" in text or "【调整" in text
                text = re.sub(r"\s*\*\*【(?:新增|调整|新增/调整)】\*\*", "", text)
                while active_levels and active_levels[-1][0] >= level:
                    active_levels.pop()
                if is_marked:
                    active_levels.append((level, True))
                active_red = any(v for _, v in active_levels)
                if level == 1 and text == "IMU 数据采集计划":
                    i += 1
                    continue
                add_heading(doc, min(level, 3), text, RED if active_red else BLUE)
            i += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, table_lines, RED if active_red else BLACK)
            continue
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run("\n".join(code_lines))
            set_run_font(run, name="Consolas", size=8.5, color=RED if active_red else BLACK)
            i += 1
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.space_after = Pt(5)
            add_text(p, stripped.lstrip("> "), RED if active_red else BLACK, bold=True, size=10)
            i += 1
            continue
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        number = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet or number:
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            p.paragraph_format.space_after = Pt(3)
            add_text(p, (bullet or number).group(1), RED if active_red else BLACK, size=10.5)
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        add_text(p, stripped, RED if active_red else BLACK, size=10.5)
        i += 1

    doc.core_properties.title = "IMU 数据采集计划（修订标红版）"
    doc.core_properties.subject = "耳机 IMU 姿势提醒与基础动作识别"
    doc.core_properties.author = "Shokz"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
