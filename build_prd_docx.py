from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "OpenSport_双设备IMU实时运动监测_PRD_V1.0.docx"
SCREENSHOT = ROOT / "assets" / "prd_dashboard.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
GRAY = "666666"
RED = "9B1C1C"
GREEN = "2E7D32"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, bold=False, color=None, font="Microsoft YaHei"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths_dxa, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(str(header)), size=font_size, bold=True, color=NAVY)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    set_run(p.add_run(text))
    return p


def add_number(doc, text, style="List Number"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    set_run(p.add_run(text))
    return p


def add_callout(doc, label, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(label + "  "), size=10.5, bold=True, color=color)
    set_run(p.add_run(text), size=10.5)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

# Running header/footer
hp = section.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(hp.add_run("OPENSPORT  ·  产品需求文档"), size=8.5, bold=True, color=GRAY)
fp = section.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(fp.add_run("内部评审稿  |  V1.0  |  2026-07-22"), size=8, color=GRAY)
# Populate even-page parts explicitly because some Word installations retain
# separate even-page header/footer relationships from the default template.
ehp = section.even_page_header.paragraphs[0]
ehp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(ehp.add_run("OPENSPORT  ·  产品需求文档"), size=8.5, bold=True, color=GRAY)
efp = section.even_page_footer.paragraphs[0]
efp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(efp.add_run("内部评审稿  |  V1.0  |  2026-07-22"), size=8, color=GRAY)

# Cover / memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
set_run(p.add_run("PRODUCT REQUIREMENTS DOCUMENT"), size=10, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
set_run(p.add_run("OpenSport 双设备 IMU 实时运动监测"), size=25, bold=True, color=NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
set_run(p.add_run("从技术 demo 到可验证产品方案"), size=13, color=GRAY)

meta = [
    ("文档状态", "内部评审稿"), ("版本", "V1.0"), ("日期", "2026-07-22"),
    ("产品阶段", "Demo 验证 / MVP 定义"), ("目标读者", "产品、设计、算法、客户端、测试、项目负责人"),
]
add_table(doc, ["字段", "内容"], meta, [1800, 7560], 9.5)

add_callout(doc, "一句话定义", "通过两台耳戴式 IMU 设备的本地实时数据与算法推理，为研发/测试人员提供运动状态、置信度、姿态和信号质量的可视化监测能力。")

add_heading(doc, "0. 评审结论摘要", 1)
add_bullet(doc, "当前 demo 已打通双设备、双数据链路、本地推理和网页实时可视化，具备继续验证的工程基础。")
add_bullet(doc, "现阶段算法仅适合作为内部探索结果：窗口级准确率 78.4%、平衡准确率 75.4%、ROC AUC 0.810；会话级准确率 84.0%。")
add_bullet(doc, "数据来自少量人员、单日、预分段动作，标签仍使用临时业务规则，因此不得对外宣称为正式产品准确率。")
add_bullet(doc, "MVP 应优先完成佩戴有效性门控、状态滞回、设备连接引导、异常提示、会话记录与可回溯日志，再扩大数据采集与跨人验证。")

add_heading(doc, "1. 背景与问题", 1)
add_heading(doc, "1.1 背景", 2)
doc.add_paragraph("OpenSport 当前已形成耳戴式 IMU 数据采集、清洗、特征提取、模型训练、实时推理和双设备监测页面的闭环。demo 支持 WitMotion CSV 记录流与 BLE GATT 直连两条隔离链路，面向研发现场展示两台设备的实时状态。")
add_heading(doc, "1.2 用户问题", 2)
add_bullet(doc, "采集人员难以快速确认两台设备是否都在稳定产出有效数据。")
add_bullet(doc, "算法人员缺少实时观察推理状态、置信度与原始信号趋势的统一入口。")
add_bullet(doc, "出现断连、WitMotion 占用、数据陈旧或佩戴异常时，当前提示不足以支持非研发用户自行排障。")
add_bullet(doc, "模型效果、业务标签与产品承诺之间尚未建立清晰边界。")

add_heading(doc, "2. 产品目标与成功指标", 1)
add_heading(doc, "2.1 产品目标", 2)
add_number(doc, "让现场人员在 30 秒内确认两台设备、数据链路与实时推理是否正常。")
add_number(doc, "让算法/测试人员可对比两侧设备的运动概率、六轴数据、姿态角与短时趋势。")
add_number(doc, "让异常状态可识别、可解释、可恢复，并保留可定位问题的会话记录。")
add_number(doc, "以内部可验证指标逐步收敛算法能力，不提前扩张为医疗、健康诊断或精准健身计量。")
add_heading(doc, "2.2 MVP 成功指标（建议）", 2)
add_table(doc, ["指标", "口径", "MVP 目标"], [
    ("双设备可用率", "有效测试时长内两台设备同时有新鲜数据的占比", "≥ 95%"),
    ("端到端可视延迟", "采样产生至页面状态更新的 P95", "记录流 ≤ 2 s；BLE ≤ 1 s"),
    ("断连可感知时间", "最后采样至页面进入异常态", "≤ 3 s"),
    ("恢复成功率", "可恢复异常中按引导恢复成功的会话占比", "≥ 90%"),
    ("内部模型基线", "按受试者分组的独立验证", "先不低于现有平衡准确率 75.4%"),
    ("状态稳定性", "稳定行为下每分钟无意义状态翻转次数", "≤ 2 次"),
], [1900, 5000, 2460], 8.8)

add_heading(doc, "3. 用户与核心场景", 1)
add_table(doc, ["角色", "核心诉求", "典型任务"], [
    ("数据采集员", "快速确认设备和数据正常", "选择链路、检查双设备、发现断连、完成采集"),
    ("算法工程师", "验证实时特征与模型输出", "观察概率、趋势、错误动作、导出日志"),
    ("测试/项目负责人", "按统一标准验收 demo", "执行测试用例、确认异常态、记录结论"),
    ("未来终端用户", "获得稳定且可理解的运动反馈", "本期不直接开放，仅作为后续产品化对象"),
], [1600, 3300, 4460], 9)

add_heading(doc, "4. 范围定义", 1)
add_table(doc, ["范围内（MVP）", "范围外（本期不做）"], [
    ("两台指定 IMU 设备的状态与实时数据展示", "任意品牌/任意数量设备的通用接入"),
    ("WitMotion 记录流与 BLE 直连二选一", "同一设备被两条链路同时连接"),
    ("运动/非运动二分类与概率展示", "动作计数、卡路里、医疗或健康诊断"),
    ("加速度、角速度、姿态角与近 180 点趋势", "长期云端历史分析与跨端账号体系"),
    ("断连、等待、重连、数据陈旧等异常提示", "自动修复第三方 WitMotion 软件问题"),
    ("本地数据与会话日志", "未经授权上传原始传感器数据"),
], [4680, 4680], 9)

add_heading(doc, "5. 当前 Demo 说明（As-Is）", 1)
if SCREENSHOT.exists():
    pic = doc.add_picture(str(SCREENSHOT), width=Inches(6.65))
    pic_paragraph = doc.paragraphs[-1]
    pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_paragraph.paragraph_format.space_after = Pt(3)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    set_run(cap.add_run("图 1  当前双设备 IMU 实时推理面板（本地 demo）"), size=8.5, color=GRAY)
add_table(doc, ["模块", "当前能力", "说明"], [
    ("链路选择", "记录流 / BLE 直连", "两条数据源隔离；同一设备不可并行占用"),
    ("设备筛选", "两台 / 单台", "仅改变页面展示，不改变采集对象"),
    ("推理", "运动 / 非运动 + 概率", "2 秒窗口、50% 重叠、阈值 0.5"),
    ("信号", "加速度、角速度、Roll/Pitch/Yaw", "展示当前值与最近 180 点趋势"),
    ("连接状态", "实时、等待、重连、陈旧", "3 秒无新样本进入 stale"),
    ("统计", "连接次数、原始帧、规范化点、源速率、错误", "用于现场诊断"),
], [1500, 2600, 5260], 8.8)

add_heading(doc, "6. 目标体验与业务流程（To-Be）", 1)
add_number(doc, "启动：用户打开监测服务，系统默认进入“WitMotion 记录流”，读取最近一次有效配置。", "List Number 2")
add_number(doc, "选择链路：若使用 WitMotion 采集，保持记录流；若使用 BLE 直连，页面明确提示先关闭 WitMotion。", "List Number 2")
add_number(doc, "设备检查：系统分别显示两台设备的连接、采样新鲜度、源速率与佩戴有效性。", "List Number 2")
add_number(doc, "推理：每台设备独立形成 2 秒窗口，完成清洗、特征提取与概率计算；满足滞回规则后切换展示状态。", "List Number 2")
add_number(doc, "监测：用户可切换单台/双台视图，观察数值与趋势；异常卡片提供原因和下一步操作。", "List Number 2")
add_number(doc, "结束：保存会话摘要、设备状态、推理序列和错误日志；原始数据遵循最小化与授权策略。", "List Number 2")

add_heading(doc, "7. 功能需求", 1)
requirements = [
    ("FR-01", "P0", "数据链路选择", "支持记录流与 BLE 直连二选一；切换后不混用历史样本。", "选择后 1 s 内更新来源说明；接口返回 source 与选择一致。"),
    ("FR-02", "P0", "双设备独立监测", "固定展示 WT22222、WT901BLE11 的独立状态和数据。", "单台异常不影响另一台刷新；设备身份与地址一致。"),
    ("FR-03", "P0", "实时推理", "展示运动/非运动、概率、最后采样时间。", "概率范围 0–100%；窗口不足时显示“正在积累窗口”。"),
    ("FR-04", "P0", "状态滞回", "连续 3 个窗口满足阈值后进入运动，连续 5 个窗口不满足后退出。", "稳定输入下不出现单窗口闪烁；规则可配置并记录版本。"),
    ("FR-05", "P0", "连接与陈旧判断", "超过 3 s 无新样本显示异常态，并区分等待、连接、重连、被占用。", "状态、颜色、文案三者一致；不只依赖颜色表达。"),
    ("FR-06", "P0", "传感器展示", "显示三轴加速度、三轴角速度、姿态角及近 180 点趋势。", "单位固定；无值用“—”而非伪造 0。"),
    ("FR-07", "P0", "佩戴有效性门控", "检测取下/不对称佩戴时暂停输出运动结论。", "页面显示“佩戴无效”；恢复稳定佩戴后再进入窗口积累。"),
    ("FR-08", "P1", "异常自助排障", "针对 WitMotion 占用、未广播、服务未启动、文件未追加给出动作指引。", "每种异常均有可执行的一步建议与重试入口。"),
    ("FR-09", "P1", "会话记录", "记录起止时间、设备、链路、模型版本、阈值、异常与摘要指标。", "可按 session_id 定位完整日志；不覆盖原始数据。"),
    ("FR-10", "P1", "设备筛选", "可查看双设备或任意单台，不改变后台采集。", "切换视图后 300 ms 内完成，后台样本计数持续增长。"),
    ("FR-11", "P2", "数据导出", "导出当前会话的摘要与推理序列。", "导出文件不默认包含个人身份字段。"),
]
add_table(doc, ["ID", "优先级", "需求", "产品规则", "验收要点"], requirements, [720, 720, 1600, 3260, 3060], 7.9)

add_heading(doc, "8. 状态、文案与异常规则", 1)
add_table(doc, ["状态", "触发条件", "用户文案", "建议动作"], [
    ("等待数据", "服务可用但无有效样本", "等待设备数据", "检查采集软件或接收器"),
    ("正在积累", "样本不足 2 秒窗口", "正在积累推理窗口", "保持设备稳定连接"),
    ("实时连接", "3 秒内持续收到样本", "实时推理中", "无需操作"),
    ("数据陈旧", "最后样本年龄 > 3 秒", "数据已暂停", "检查链路或刷新"),
    ("正在重连", "接收器执行重连", "设备重连中", "等待；超时后展示原因"),
    ("WitMotion 占用", "BLE 无法连接且检测到占用可能", "设备可能被 WitMotion 占用", "关闭 WitMotion 后重试"),
    ("佩戴无效", "取下或不对称佩戴门控命中", "请重新佩戴设备", "调整后静止 5–10 秒校准"),
    ("服务不可用", "API 请求失败", "本地监控服务未启动", "启动服务并重新加载"),
], [1500, 2900, 2300, 2660], 8.5)

add_heading(doc, "9. 数据与算法要求", 1)
add_heading(doc, "9.1 当前基线", 2)
add_table(doc, ["项目", "当前值", "产品解释"], [
    ("模型", "L2 正则化逻辑回归", "解释性较好，作为 demo 基线"),
    ("样本", "97 个文件、50 个采集会话、7,638 个窗口", "样本规模有限"),
    ("窗口", "2 秒、50% 重叠", "约每 1 秒产生一次新推理"),
    ("阈值", "0.5", "尚未按业务成本系统校准"),
    ("窗口级", "准确率 78.4%；平衡准确率 75.4%；AUC 0.810", "仅内部嵌套分组验证"),
    ("会话级", "准确率 84.0%；平衡准确率 85.4%；AUC 0.884", "窗口概率均值后的结果"),
], [1700, 3500, 4160], 8.8)
add_callout(doc, "重要边界", "上述结果来自少量人员、单日和预先分段的数据。标签中的“运动”实际按临时的 exercise/non_exercise 协议映射，ambiguous 与 wear_artifact 未参与训练；不能推导新用户、自由生活或坐姿健康判断的真实准确率。", RED)
add_heading(doc, "9.2 产品化数据要求", 2)
add_bullet(doc, "每段必须保存 subject_id、device_id、session_id、raw_action、fitness_state、motion_state、posture_state、wear_state、context、采样率、设备方向、起止时间和备注。")
add_bullet(doc, "每人每动作至少 3–5 段、每段 1–3 分钟，覆盖不同人员、日期、松紧度与真实负样本；验证按人划分。")
add_bullet(doc, "每次佩戴静止 5–10 秒完成重力与零偏校准；保存原始数据，清洗结果另存并保留日志。")
add_bullet(doc, "模型包必须包含模型版本、特征版本、窗口参数、阈值、训练数据版本和评估报告。")

add_heading(doc, "10. 非功能需求", 1)
add_table(doc, ["类别", "要求"], [
    ("性能", "页面轮询与渲染不阻塞；双设备持续运行 2 小时无明显内存增长；P95 延迟符合第 2.2 节。"),
    ("可靠性", "单设备断连不影响另一设备；CSV 追加读取不得重复导入同一会话；重启后状态可恢复。"),
    ("安全与隐私", "默认本地处理；原始传感器和身份信息不提交代码仓库、不自动上传；导出遵循最小化原则。"),
    ("可访问性", "状态同时使用文字、图标与颜色；控件支持键盘操作；正文与背景对比度达到 WCAG AA。"),
    ("兼容性", "MVP 支持 Windows 10/11 与最新版 Edge/Chrome；窄屏采用单列卡片。"),
    ("可观测性", "记录连接、帧数、规范化样本、源速率、丢弃字节、最后错误、模型与配置版本。"),
], [1800, 7560], 9)

add_heading(doc, "11. 埋点与评估方案", 1)
add_table(doc, ["事件/指标", "关键属性", "用途"], [
    ("session_start / end", "session_id、source、model_version、device_count", "计算会话成功率与时长"),
    ("source_switch", "from、to、result、duration_ms", "评估链路切换体验"),
    ("device_state_change", "device_id、from、to、reason、age_ms", "定位断连与状态抖动"),
    ("inference_change", "device_id、label、probability、threshold", "评估输出稳定性"),
    ("recovery_action", "error_type、action、result", "优化异常指引"),
    ("export_result", "format、row_count、contains_raw", "审计数据导出"),
], [2200, 4400, 2760], 8.8)

add_heading(doc, "12. Figma 原型与视觉交付规范", 1)
doc.add_paragraph("设计稿建议以当前深色监测台为视觉基线，但交付时必须从“技术展示页”升级为“任务导向的状态监测工具”。")
add_heading(doc, "12.1 页面与组件", 2)
add_bullet(doc, "Frame：Desktop 1440、Desktop 1280、Tablet 768、Mobile 375；桌面双列，≤1050 px 单列。")
add_bullet(doc, "页面：首页/监测台、链路切换确认、设备异常、佩戴校准、会话结束摘要、空状态。")
add_bullet(doc, "组件：Source Select、Device Filter、Device Card、Connection Chip、Inference State、Probability Bar、Sensor Tile、Trend Chart、Error Banner、Action Button。")
add_bullet(doc, "组件变量：deviceState（waiting/live/stale/reconnecting/error）、inference（unknown/active/idle/invalidWear）、density（desktop/mobile）。")
add_heading(doc, "12.2 标注要求", 2)
add_table(doc, ["项目", "交付标准"], [
    ("布局", "8 px 基础栅格；卡片内边距 16/24；桌面内容最大宽度 1680；断点和重排规则写入 Dev Mode。"),
    ("颜色", "建立 Background/Surface/Border/Text/Info/Success/Warning/Danger token；不得只用颜色区分状态。"),
    ("文字", "中文优先；状态文案与第 8 节一致；数值单位固定，不重复混排。"),
    ("交互", "选择、悬停、聚焦、禁用、加载、错误、空态均需组件状态；切换链路需说明占用风险。"),
    ("图表", "标注时间窗、采样点数、坐标范围、自适应规则、缺失点处理和图例颜色。"),
    ("开发交付", "图层语义命名；组件与变量可复用；关键流程连线；需求 ID 写入注释；图标使用 SVG。"),
], [1800, 7560], 9)
add_callout(doc, "Figma 评审重点", "设计评审不能只看正常态。至少用 8 个状态场景走查：双设备正常、单设备断连、两台均无数据、窗口积累、WitMotion 占用、BLE 重连、佩戴无效、API 不可用。")

add_heading(doc, "13. 验收测试清单", 1)
tests = [
    ("AC-01", "记录流正常追加", "两台卡片独立更新；source=witmonitor；无 BLE 占用提示"),
    ("AC-02", "切换 BLE 直连", "展示关闭 WitMotion 提示；来源切换后不混入记录流样本"),
    ("AC-03", "单台断连 >3 秒", "该卡片进入数据陈旧/重连；另一台继续刷新"),
    ("AC-04", "窗口不足", "显示正在积累，不显示运动概率 0%"),
    ("AC-05", "单窗口越过阈值", "不立即切换状态；满足连续窗口规则后再切换"),
    ("AC-06", "检测到佩戴异常", "暂停运动结论并引导重新佩戴与校准"),
    ("AC-07", "API 停止", "页面显示服务不可用；恢复后自动继续，无需刷新整页"),
    ("AC-08", "窄屏 375 px", "卡片单列、无横向溢出、控件可操作"),
    ("AC-09", "连续运行 2 小时", "无明显内存泄漏；计数单调；页面仍可交互"),
    ("AC-10", "结束会话并导出", "文件包含 session、设备、模型与异常摘要；无未授权个人信息"),
]
add_table(doc, ["编号", "场景", "预期结果"], tests, [1000, 3100, 5260], 8.8)

add_heading(doc, "14. 版本路线图", 1)
add_table(doc, ["阶段", "目标", "核心交付"], [
    ("V1.0 Demo 固化", "可重复演示与测试", "状态机、滞回、佩戴门控、异常文案、测试用例"),
    ("V1.1 内部试用", "提升稳定性与可追溯性", "会话管理、导出、日志、性能与长稳测试"),
    ("V1.2 算法验证", "获得可信的跨人结果", "扩大数据、多日多人与真实负样本、按人独立验证"),
    ("V2.0 产品化探索", "面向真实用户验证价值", "用户研究、任务闭环、隐私机制、设备与客户端集成"),
], [1600, 2600, 5160], 9)

add_heading(doc, "15. 风险与待决策项", 1)
add_table(doc, ["类型", "问题", "建议决策"], [
    ("业务定义", "“运动”是人体活动、健身状态还是具体动作？", "V1 明确为协议定义的健身/非健身；文案避免泛化"),
    ("算法", "开合跳、弯腰取物、跑步机等存在明显混淆", "补数据与重做标签前，不扩大能力承诺"),
    ("佩戴", "设备取下/松动可能造成强信号并被误判", "佩戴有效性作为推理前置门控"),
    ("链路", "WitMotion 与 BLE 可能互相占用", "产品交互坚持二选一并提供切换引导"),
    ("姿态", "绝对角度跨人、跨佩戴不可直接比较", "每次佩戴做中立姿态校准，使用相对角度"),
    ("隐私", "原始 IMU 与人员标签可能构成敏感研究数据", "默认本地、最小采集、分离身份映射并控制导出"),
], [1400, 3980, 3980], 8.8)

add_heading(doc, "附录 A：需求评审检查表", 1)
for item in [
    "业务目标、目标用户和使用场景已由负责人确认。",
    "“运动/非运动”的标签定义与对外文案一致。",
    "P0 功能、异常态、埋点与验收用例已完成评审。",
    "Figma 组件状态、响应式断点和开发标注已齐全。",
    "算法报告包含按人划分验证、阈值依据和失败案例。",
    "隐私、数据留存、导出范围和权限边界已确认。",
    "发布前完成双设备长稳、断连恢复和链路切换测试。",
]:
    add_bullet(doc, "□ " + item)

doc.core_properties.title = "OpenSport 双设备 IMU 实时运动监测 PRD"
doc.core_properties.subject = "产品需求文档"
doc.core_properties.author = "OpenSport 产品团队"
doc.core_properties.keywords = "PRD, IMU, 实时推理, 双设备, OpenSport"
doc.save(OUT)
print(OUT)
